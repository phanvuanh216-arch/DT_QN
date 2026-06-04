"""
modules/export_ban_tin.py
Module EXPORT BẢN TIN – xuất bản tin ra file .docx theo định dạng mẫu
"""
import streamlit as st
import json
import os
import tempfile
from datetime import datetime
from utils.data_fetcher import DOI_TUONG, KY_THANG, RISK_LABELS, XA_LIST, load_data_for_xa


def render():
    st.markdown('<div class="module-header">📤 MODULE: Export bản tin</div>', unsafe_allow_html=True)

    ban_tin_list = st.session_state.get("ban_tin_list", [])

    st.markdown("### Chọn bản tin cần xuất")

    export_mode = st.radio(
        "Chế độ export:",
        ["📄 Xuất từ bản tin đã lưu", "⚡ Xuất trực tiếp (chọn xã + đối tượng)"],
        horizontal=True,
    )

    if export_mode == "📄 Xuất từ bản tin đã lưu":
        _export_from_saved(ban_tin_list)
    else:
        _export_direct()


def _export_from_saved(ban_tin_list):
    if not ban_tin_list:
        st.warning("Chưa có bản tin nào. Hãy tạo bản tin ở module **Truyền thông xã** trước.")
        return

    options = {f"[{bt['created_at']}] Xã {bt['xa']} – {bt['doi_tuong_name']}": bt for bt in ban_tin_list}
    chon = st.multiselect("Chọn bản tin cần xuất:", list(options.keys()))

    if not chon:
        return

    col1, col2 = st.columns(2)
    with col1:
        if st.button("📤 Export DOCX", type="primary", use_container_width=True):
            selected = [options[k] for k in chon]
            for bt in selected:
                _do_export_docx(bt)
    with col2:
        if st.button("📊 Export JSON (dữ liệu thô)", use_container_width=True):
            selected = [options[k] for k in chon]
            for bt in selected:
                _do_export_json(bt)


def _export_direct():
    col1, col2 = st.columns(2)
    with col1:
        xa_chon = st.selectbox("🏘️ Xã:", XA_LIST, key="exp_xa")
    with col2:
        dt_chon = st.selectbox("🌱 Đối tượng:", ["Lúa", "Rau", "Lợn", "Gà"], key="exp_dt")

    dt_map = {"Lúa": "lua", "Rau": "rau", "Lợn": "lon", "Gà": "ga"}
    dt_key = dt_map[dt_chon]

    if st.button("⚡ Tải & Export ngay", type="primary"):
        with st.spinner("Đang tải dữ liệu và tạo bản tin..."):
            data, source = load_data_for_xa(xa_chon)
            dt_info = DOI_TUONG[dt_key]
            rr_data = data.get("rui_ro", {}).get(dt_key, {})
            from modules.ban_tin_xa import _generate_nhan_dinh
            nhan_dinh = _generate_nhan_dinh(xa_chon, dt_chon, rr_data)

            bt = {
                "id": f"{xa_chon}_{dt_key}_direct",
                "xa": xa_chon,
                "doi_tuong_key": dt_key,
                "doi_tuong_name": dt_chon,
                "nhan_dinh": nhan_dinh,
                "rui_ro_data": rr_data,
                "du_bao": data.get("du_bao_thang", {}),
                "created_at": datetime.now().strftime("%d/%m/%Y %H:%M"),
            }
            _do_export_docx(bt)


def _do_export_docx(bt: dict):
    """Tạo file DOCX theo mẫu bản tin cảnh báo khí hậu."""
    try:
        docx_bytes = _build_docx(bt)
        filename = f"BanTin_{bt['xa'].replace(' ', '_')}_{bt['doi_tuong_name']}_{datetime.now().strftime('%Y%m%d')}.docx"
        st.download_button(
            label=f"⬇️ Tải xuống: {filename}",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl_{bt['id']}_{datetime.now().microsecond}",
        )
        st.success(f"✅ Đã tạo bản tin DOCX cho **{bt['doi_tuong_name']}** – xã **{bt['xa']}**")
    except Exception as e:
        st.error(f"❌ Lỗi tạo DOCX: {str(e)}")
        st.info("💡 Gợi ý: Cài thư viện `pip install python-docx` trên môi trường Streamlit.")


def _do_export_json(bt: dict):
    json_str = json.dumps(bt, ensure_ascii=False, indent=2)
    filename = f"BanTin_{bt['xa'].replace(' ', '_')}_{bt['doi_tuong_name']}.json"
    st.download_button(
        label=f"⬇️ Tải JSON: {filename}",
        data=json_str.encode("utf-8"),
        file_name=filename,
        mime="application/json",
        key=f"json_{bt['id']}_{datetime.now().microsecond}",
    )


def _build_docx(bt: dict) -> bytes:
    """Xây dựng file DOCX từ dữ liệu bản tin (dùng python-docx)."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    doc = Document()

    # Thiết lập trang A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    def set_cell_bg(cell, hex_color):
        """Set cell background color."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    RISK_HEX = {0: "FFFFFF", 1: "FFFDE7", 2: "FFE0B2", 3: "FFCDD2"}
    RISK_TEXT = {0: "Không", 1: "Cấp 1", 2: "Cấp 2", 3: "Cấp 3"}

    xa = bt["xa"]
    doi_tuong = bt["doi_tuong_name"]
    nhan_dinh = bt["nhan_dinh"]
    du_bao = bt.get("du_bao", {})
    rr_data = bt.get("rui_ro_data", {})
    dt_key = bt.get("doi_tuong_key", "")
    dt_info = DOI_TUONG.get(dt_key, {})
    created_at = bt.get("created_at", "")

    # ── Tiêu đề ───────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"BẢN TIN CẢNH BÁO KHÍ HẬU THÁNG 6-8 KHU VỰC XÃ {xa.upper()}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    # Khu vực
    doc.add_paragraph()
    tbl_header = doc.add_table(rows=2, cols=2)
    tbl_header.style = "Table Grid"
    tbl_header.cell(0, 0).text = "Bản tin khí hậu cho khu vực"
    tbl_header.cell(0, 1).text = ""
    tbl_header.cell(1, 0).text = "Xã"
    c11 = tbl_header.cell(1, 1)
    run = c11.paragraphs[0].add_run(xa)
    run.bold = True
    set_cell_bg(tbl_header.cell(0, 0), "BDD7EE")
    set_cell_bg(tbl_header.cell(1, 0), "D9E1F2")

    doc.add_paragraph()

    # ── Dự báo khí hậu ────────────────────────────────────────────────────────
    h2 = doc.add_paragraph()
    r = h2.add_run("Dự báo khí hậu tháng 6 đến tháng 8/2026")
    r.bold = True
    r.font.size = Pt(12)

    # Bảng dự báo (3 hàng: header + nhiệt độ + lượng mưa, 10 cột)
    cols_db = 1 + len(KY_THANG)  # "Chỉ tiêu" + 9 kỳ
    tbl_db = doc.add_table(rows=4, cols=cols_db)
    tbl_db.style = "Table Grid"

    # Hàng 0: header
    tbl_db.cell(0, 0).text = "Dự báo tháng"
    for j, ky in enumerate(KY_THANG):
        tbl_db.cell(0, j + 1).text = ky
        set_cell_bg(tbl_db.cell(0, j + 1), "FFF2CC")

    set_cell_bg(tbl_db.cell(0, 0), "FFF2CC")

    # Hàng 1-2: Tỷ lệ (Thấp/Xấp xỉ/Cao) – compact
    tbl_db.cell(1, 0).text = "Nhiệt độ TBNN (%)\nThấp / Xấp xỉ / Cao"
    tbl_db.cell(2, 0).text = "Lượng mưa TBNN (%)\nThấp / Xấp xỉ / Cao"
    tbl_db.cell(3, 0).text = "Nhận định"

    for j, ky in enumerate(KY_THANG):
        d = du_bao.get(ky, {})
        nt = d.get("nhiet_do", {})
        lm = d.get("luong_mua", {})
        nt_str = f"{nt.get('thap_hon','')} / {nt.get('xap_xi','')} / {nt.get('cao_hon','')}"
        lm_str = f"{lm.get('thap_hon','')} / {lm.get('xap_xi','')} / {lm.get('cao_hon','')}"
        tbl_db.cell(1, j + 1).text = nt_str
        tbl_db.cell(2, j + 1).text = lm_str
        nhan_dinh_ky = d.get("nhiet_do", {}).get("nhan_dinh", "")
        tbl_db.cell(3, j + 1).text = ""

    set_cell_bg(tbl_db.cell(1, 0), "DEEAF1")
    set_cell_bg(tbl_db.cell(2, 0), "E2EFDA")

    doc.add_paragraph()

    # ── Bảng rủi ro theo đối tượng ────────────────────────────────────────────
    def add_risk_table(doc, title, chu_ky, rui_ro_kh, dich_hai_list, rr_dict):
        h3 = doc.add_paragraph()
        r = h3.add_run(title)
        r.bold = True
        r.font.size = Pt(11)

        all_rows = []
        if chu_ky:
            all_rows.append(("Chu kỳ sinh trưởng", "CHU_KY", chu_ky))
        for rr in rui_ro_kh:
            all_rows.append((rr, "KHI_HAU", rr_dict.get(rr, {})))
        if dich_hai_list:
            for dh in dich_hai_list:
                all_rows.append((dh, "DICH_HAI", rr_dict.get(dh, {})))

        n_rows = 1 + len(all_rows)
        n_cols = 1 + len(KY_THANG)
        tbl = doc.add_table(rows=n_rows, cols=n_cols)
        tbl.style = "Table Grid"

        # Header row
        tbl.cell(0, 0).text = "Giai đoạn"
        set_cell_bg(tbl.cell(0, 0), "1F3864")
        tbl.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tbl.cell(0, 0).paragraphs[0].runs[0].bold = True
        for j, ky in enumerate(KY_THANG):
            c = tbl.cell(0, j + 1)
            c.text = ky
            set_cell_bg(c, "1F3864")
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                c.paragraphs[0].runs[0].bold = True

        # Data rows
        for i, (indicator, row_type, ky_vals) in enumerate(all_rows):
            row_idx = i + 1
            tbl.cell(row_idx, 0).text = indicator

            if row_type == "CHU_KY":
                set_cell_bg(tbl.cell(row_idx, 0), "FCE4D6")
                for j, ky in enumerate(KY_THANG):
                    val = ky_vals.get(ky, "")
                    tbl.cell(row_idx, j + 1).text = str(val)
                    set_cell_bg(tbl.cell(row_idx, j + 1), "FCE4D6")
            elif row_type == "KHI_HAU":
                set_cell_bg(tbl.cell(row_idx, 0), "DEEAF1")
                for j, ky in enumerate(KY_THANG):
                    val = ky_vals.get(ky, 0) if isinstance(ky_vals, dict) else 0
                    c = tbl.cell(row_idx, j + 1)
                    c.text = RISK_TEXT.get(val, "")
                    set_cell_bg(c, RISK_HEX.get(val, "FFFFFF"))
            else:  # DICH_HAI
                for j, ky in enumerate(KY_THANG):
                    val = ky_vals.get(ky, 0) if isinstance(ky_vals, dict) else 0
                    c = tbl.cell(row_idx, j + 1)
                    c.text = RISK_TEXT.get(val, "")
                    set_cell_bg(c, RISK_HEX.get(val, "FFFFFF"))

        doc.add_paragraph()

    # Thêm bảng rủi ro
    if dt_key == "rau":
        for loai_key, loai_info in dt_info.get("loai_rau", {}).items():
            loai_rr = rr_data.get(loai_key, {})
            add_risk_table(
                doc,
                f"Mức độ rủi ro đối với cây {loai_info['name']} trong giai đoạn Tháng 6 đến Tháng 8 năm 2026",
                loai_info.get("chu_ky", {}),
                dt_info.get("rui_ro_khi_hau", []),
                loai_info.get("dich_hai", []),
                loai_rr,
            )
    else:
        entity_word = "cây" if dt_key == "lua" else "con"
        add_risk_table(
            doc,
            f"Mức độ rủi ro đối với {entity_word} {doi_tuong} trong giai đoạn Tháng 6 đến Tháng 8 năm 2026",
            dt_info.get("chu_ky", {}),
            dt_info.get("rui_ro_khi_hau", []),
            dt_info.get("dich_hai", dt_info.get("dich_benh", [])),
            rr_data,
        )

    # ── Nhận định ─────────────────────────────────────────────────────────────
    h_nd = doc.add_paragraph()
    r = h_nd.add_run("Nhận định và khuyến cáo:")
    r.bold = True
    r.font.size = Pt(12)

    for line in nhan_dinh.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph(f"Bản tin tạo lúc: {created_at} | Nguồn: http://222.254.32.10/forecast/Detai_QuangNinh/")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Lưu vào bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
